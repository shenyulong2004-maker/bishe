#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
生成家校合作平台系统Word文档
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


def shade_cell(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)


def set_table_header(row, headers, bg_color='1F3864'):
    for i, header in enumerate(headers):
        cell = row.cells[i]
        shade_cell(cell, bg_color)
        para = cell.paragraphs[0]
        run = para.add_run(header)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(11)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_document():
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)

    # ===================== 封面 =====================
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('家校合作平台系统', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 56, 100)
        run.font.name = '黑体'

    sub = doc.add_paragraph('系 统 设 计 文 档')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.runs[0]
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(68, 114, 196)
    sub_run.font.name = '黑体'

    doc.add_paragraph()

    info = doc.add_paragraph('版本：V1.0      日期：2026年3月22日')
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.runs[0].font.size = Pt(12)
    info.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ===================== 1. 系统功能 =====================
    h1 = doc.add_heading('1. 系统功能', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(31, 56, 100)

    doc.add_paragraph('家校合作平台面向家长、教师和管理员三类用户，提供全方位的家校互动与管理功能。')

    # 功能表格
    func_table = doc.add_table(rows=1, cols=3)
    func_table.style = 'Table Grid'
    set_table_header(func_table.rows[0], ['家长端功能', '教师端功能', '管理员功能'])

    func_data = [
        ('学生成绩查询与分析', '学生成绩录入与管理', '用户管理与权限控制'),
        ('学生考勤信息实时推送', '考勤管理与统计', '班级与课程管理'),
        ('与教师在线沟通反馈', '作业发布与在线批改', '系统数据统计分析'),
        ('学生作业提交与查看', '班级通知与公告发布', '平台公告与系统维护'),
        ('班级动态与通知公告', '学生学习数据分析', '用户反馈处理'),
        ('学生学习进度跟踪', '教学资源共享库', '系统日志与安全审计'),
        ('家庭作业辅导资源', '家校沟通记录管理', '数据报表与导出'),
    ]

    for row_data in func_data:
        row = func_table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ===================== 2. 系统架构 =====================
    h1 = doc.add_heading('2. 系统架构', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(31, 56, 100)

    doc.add_paragraph('系统采用经典的五层分层架构，各层职责清晰、松耦合、易扩展。')

    # 分层架构表格
    arch_table = doc.add_table(rows=1, cols=3)
    arch_table.style = 'Table Grid'
    set_table_header(arch_table.rows[0], ['架构层级', '层级说明', '核心技术组件'])

    arch_data = [
        ('表现层（前端）', '负责用户交互与界面展示', 'Vue 3 / TypeScript / Vite / Element Plus'),
        ('接口层（API）', '提供RESTful接口，处理请求路由', 'Express.js / FastAPI / Nginx'),
        ('业务层（Service）', '核心业务逻辑、数据验证、权限控制', 'Spring Boot / Node.js 业务模块'),
        ('数据访问层（DAO）', '数据库操作封装，屏蔽底层差异', 'Sequelize / SQLAlchemy / MyBatis'),
        ('存储层（Storage）', '数据持久化、缓存与文件存储', 'MySQL 8.0 / Redis / OSS'),
    ]

    for row_data in arch_data:
        row = arch_table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # 2.1 后端架构
    h2 = doc.add_heading('2.1 后端架构', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(68, 114, 196)

    backend_items = [
        '微服务划分：用户服务、成绩服务、考勤服务、通知服务、作业服务独立部署',
        '消息队列：RabbitMQ 实现异步消息处理，削峰填谷，提升系统吞吐量',
        '缓存层：Redis 6.0 用于会话管理、热点数据缓存，降低数据库压力',
        '日志系统：ELK Stack（Elasticsearch + Logstash + Kibana）集中收集与分析日志',
        '监控告警：Prometheus + Grafana 实时监控系统健康状态与性能指标',
        '容器化部署：Docker + Kubernetes 实现弹性伸缩与高可用',
    ]
    for item in backend_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item).font.size = Pt(11)

    doc.add_paragraph()

    # 2.2 数据库架构
    h2 = doc.add_heading('2.2 数据库架构', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(68, 114, 196)

    db_table = doc.add_table(rows=1, cols=4)
    db_table.style = 'Table Grid'
    set_table_header(db_table.rows[0], ['数据库', '类型', '用途', '主要存储内容'])

    db_data = [
        ('MySQL 8.0', '关系型数据库', '主数据存储', '用户、班级、成绩、考勤、作业等核心数据'),
        ('Redis 6.0', '内存数据库', '缓存与会话', '登录Token、热点数据、消息队列'),
        ('MongoDB', '文档数据库', '非结构化存储', '操作日志、用户行为记录'),
        ('Elasticsearch', '搜索引擎', '全文检索', '资料搜索、日志分析'),
    ]

    for row_data in db_data:
        row = db_table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ===================== 3. 技术栈 =====================
    h1 = doc.add_heading('3. 技术栈', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(31, 56, 100)

    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = 'Table Grid'
    set_table_header(tech_table.rows[0], ['技术分类', '技术选型', '说明'])

    tech_data = [
        # 前端
        ('前端框架', 'Vue 3 + TypeScript', '组合式API，类型安全，开发效率高'),
        ('构建工具', 'Vite 5', '极速热更新，模块化构建'),
        ('状态管理', 'Pinia', '轻量级，Vue3原生支持'),
        ('UI组件库', 'Element Plus', '功能完善，适合管理平台'),
        ('图表可视化', 'ECharts 5', '丰富图表类型，数据分析展示'),
        ('实时通信', 'Socket.io', 'WebSocket封装，消息实时推送'),
        # 样式
        ('CSS预处理', 'SCSS', '变量/嵌套/混入，样式工程化'),
        ('CSS框架', 'Tailwind CSS', '原子化CSS，快速构建响应式页面'),
        ('动画库', 'Animate.css / GSAP', '页面交互动画效果'),
        # 后端
        ('后端语言', 'Node.js / Python', '高性能异步处理与数据科学支持'),
        ('后端框架', 'Express.js / FastAPI', 'RESTful API开发，自动生成文档'),
        ('身份认证', 'JWT + OAuth 2.0', '无状态认证，支持第三方登录'),
        ('API文档', 'Swagger / OpenAPI', '自动生成接口文档，便于联调'),
        # 存储
        ('主数据库', 'MySQL 8.0', '事务支持，数据一致性保障'),
        ('缓存数据库', 'Redis 6.0', '高性能键值存储，毫秒级响应'),
        ('文件存储', '阿里云OSS', '图片/文件分布式存储'),
        ('消息队列', 'RabbitMQ', '异步解耦，削峰填谷'),
        # 运维
        ('容器化', 'Docker + K8s', '一致性部署，弹性伸缩'),
        ('CI/CD', 'GitHub Actions', '自动化测试与部署'),
        ('监控', 'Prometheus + Grafana', '实时监控，可视化告警'),
    ]

    for row_data in tech_data:
        row = tech_table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ===================== 4. 创新点 =====================
    h1 = doc.add_heading('4. 系统创新点', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(31, 56, 100)

    # 4.1 交互优化
    h2 = doc.add_heading('4.1 交互优化', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(68, 114, 196)

    inter_table = doc.add_table(rows=1, cols=2)
    inter_table.style = 'Table Grid'
    set_table_header(inter_table.rows[0], ['创新特性', '实现方案'])

    inter_data = [
        ('智能消息推送', '基于WebSocket实现秒级实时通知，成绩、考勤变动即时推送至家长'),
        ('个性化仪表板', '拖拽式组件布局，用户可自定义关注数据模块'),
        ('深色/浅色主题', 'CSS变量驱动主题系统，一键切换，保护视力'),
        ('渐进式加载', '骨架屏+懒加载，首屏加载时间压缩至1秒内'),
        ('离线缓存', 'Service Worker缓存策略，弱网环境下可查看历史数据'),
        ('无障碍设计', '遵循WCAG 2.1 AA标准，支持键盘导航与屏幕阅读器'),
    ]

    for row_data in inter_data:
        row = inter_table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # 4.2 题型灵活
    h2 = doc.add_heading('4.2 题型灵活', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(68, 114, 196)

    question_items = [
        '多元题型支持：单选题、多选题、填空题、简答题、文件上传题、连线题、排序题',
        '动态题库管理：支持按知识点、难度等级、年级学科多维度分类管理',
        '智能自动组卷：基于知识点覆盖率与难度分布算法，一键生成平衡试卷',
        '富文本编辑器：题目支持公式（MathJax）、图片、音频、视频等富媒体内容',
        '批量导入导出：兼容Excel/Word模板批量导入题目，支持一键导出试卷',
        '题目版本控制：记录每次修改历史，支持版本对比与回滚',
        '错题本功能：自动汇总学生错题，针对性生成练习集',
    ]
    for item in question_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item).font.size = Pt(11)

    doc.add_paragraph()

    # 4.3 数据可靠
    h2 = doc.add_heading('4.3 数据可靠', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(68, 114, 196)

    reliable_table = doc.add_table(rows=1, cols=2)
    reliable_table.style = 'Table Grid'
    set_table_header(reliable_table.rows[0], ['可靠性维度', '技术实现'])

    reliable_data = [
        ('数据加密', '传输层TLS 1.3，敏感字段AES-256加密存储'),
        ('自动备份', '每日全量备份 + 实时增量备份，多地域冗余存储'),
        ('时间点恢复', '支持任意时间点数据恢复（PITR），最大限度减少数据丢失'),
        ('审计日志', '全量操作审计追踪，记录用户每次数据变更'),
        ('权限隔离', '基于RBAC细粒度权限控制，最小权限原则'),
        ('高可用架构', 'MySQL主从复制+读写分离，Redis哨兵模式，故障自动切换'),
        ('数据一致性', '分布式事务（Saga模式），保障跨服务数据最终一致性'),
    ]

    for row_data in reliable_data:
        row = reliable_table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ===================== 5. 性能指标 =====================
    h1 = doc.add_heading('5. 性能指标', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(31, 56, 100)

    perf_table = doc.add_table(rows=1, cols=3)
    perf_table.style = 'Table Grid'
    set_table_header(perf_table.rows[0], ['性能指标', '目标值', '实现手段'])

    perf_data = [
        ('首页加载时间', '< 2 秒', '代码分割 + CDN加速 + Gzip压缩'),
        ('API响应时间', '< 500 ms', 'Redis缓存 + 数据库索引优化'),
        ('系统可用性', '>= 99.9%', '多节点部署 + 健康检查 + 自动重启'),
        ('支持并发用户', '>= 10,000', 'Kubernetes弹性扩容 + 负载均衡'),
        ('数据库查询', '< 100 ms', '索引优化 + 读写分离 + 连接池'),
        ('消息推送延迟', '< 1 秒', 'WebSocket长连接 + 消息队列'),
    ]

    for row_data in perf_data:
        row = perf_table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = text
            cell.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ===================== 6. 安全性设计 =====================
    h1 = doc.add_heading('6. 安全性设计', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor(31, 56, 100)

    security_items = [
        '身份认证：JWT Token + 多因素认证（MFA），防止账号盗用',
        '防SQL注入：ORM参数化查询，严格校验用户输入',
        '防XSS攻击：前端输入过滤 + CSP内容安全策略',
        '防CSRF攻击：Token验证机制 + SameSite Cookie策略',
        'DDoS防护：Nginx限流 + 云WAF防火墙',
        '接口加签：敏感接口请求签名验证，防止篡改',
        '数据脱敏：手机号、身份证等敏感信息展示脱敏',
        '合规遵从：符合《个人信息保护法》《网络安全法》相关要求',
    ]
    for item in security_items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(item).font.size = Pt(11)

    doc.add_paragraph()

    # ===================== 页脚 =====================
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run('--- 家校合作平台系统设计文档  v1.0  |  2026年3月22日 ---')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 150, 150)

    # 保存
    output_path = r'd:\bishe\家校合作平台系统设计文档.docx'
    doc.save(output_path)
    print('文档已成功生成：' + output_path)


if __name__ == '__main__':
    create_document()
